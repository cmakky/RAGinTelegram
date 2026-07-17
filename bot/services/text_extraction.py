"""
Извлечение чистого текста из файлов поддерживаемых форматов.
"""
import logging
from pathlib import Path

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

# Кодировки перебираются по очереди
TEXT_ENCODING = ("utf-8", "cp1251", "latin-1")


class ExtractionError(Exception):
    "Не удалось извлечь текст из файла."


def extract_text(file_path: Path) -> str:
    """
    Извлекает текст из файла по его расширению.
    Исключение если формат не поддерживается или
    результат оказался пустым
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf(file_path)
    elif suffix == ".docx":
        text = _extract_docx(file_path)
    elif suffix in (".txt", ".md"):
        text = _extract_plain_text(file_path)
    else:
        raise ExtractionError(f'Формат {suffix} не поддерживается.')

    text = _normalize_whitespace(text)

    if not text.strip():
        raise ExtractionError(f'Не удалось извлечь текст из файла.')
    
    return text


def _extract_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages_text = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text:
            pages_text.append(page_text)
        else:
            logger.warning("Страница %d в %s не содержит текста", i + 1, file_path.name)

    return "\n\n".join(pages_text)


def _extract_docx(file_path: Path) -> str:
    document = Document(str(file_path))
    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            if any(cells_text):
                parts.append(" | ".join(cells_text))
    
    return "\n".join(parts)


def _extract_plain_text(file_path: Path) -> str:
    raw_bytes = file_path.read_bytes()

    for encoding in TEXT_ENCODING:
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    
    return raw_bytes.decode("utf-8", errors="replace")


def _normalize_whitespace(text: str) -> str:
    """Убирает пустые символы и пробелы по краям"""
    lines = [line.strip() for line in text.splitlines()]
    cleaned_lines = []
    previous_blank = False

    for line in lines:
        is_blank = line == ""
        if is_blank and previous_blank:
            continue
        cleaned_lines.append(line)
        previous_blank = is_blank

    return "\n".join(cleaned_lines).strip()
